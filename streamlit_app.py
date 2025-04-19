import streamlit as st
import pandas as pd
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import re
from io import BytesIO
import os

st.set_page_config(page_title="Calificador de Exámenes", layout="wide")
st.title("📄 Sistema Automático de Corrección de Exámenes en Word")

# Inicializar clave dinámica para limpiar archivos
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = "alumno_upload_1"

# Función para reiniciar clave de uploader (limpiar archivos)
def limpiar_archivos_alumnos():
    nueva_clave = "alumno_upload_" + str(int(st.session_state.uploader_key.split("_")[-1]) + 1)
    st.session_state.uploader_key = nueva_clave

# --- Subida de archivos ---
st.markdown("Sube el archivo `.docx` con la **clave correcta** (extraerá exactamente 10 respuestas de la tabla).")
clave_file = st.file_uploader("📥 Sube el archivo .docx con la clave correcta", type=["docx"], key="clave")

st.markdown("Sube uno o varios archivos `.docx` de los **alumnos**. Se evaluará resaltado, subrayado y tabla (ignorando mayúsculas/minúsculas).")
uploaded_files = st.file_uploader(
    "📥 Sube uno o varios archivos .docx de los alumnos",
    type=["docx"],
    accept_multiple_files=True,
    key=st.session_state.uploader_key,
    label_visibility="collapsed"
)

if uploaded_files:
    st.button("🧹 Quitar todos los archivos", on_click=limpiar_archivos_alumnos)

# --- Funciones ---
def extraer_clave_de_tabla(document):
    respuestas = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                letra = cell.text.strip().upper()
                if letra in ['A', 'B', 'C', 'D']:
                    respuestas.append(letra)
                if len(respuestas) == 10:
                    return respuestas
    return respuestas[:10]

def extraer_nombre(document, filename):
    for para in document.paragraphs:
        match = re.search(r"NOMBRE[S]?:?\s*(.*)", para.text.upper())
        if match:
            return match.group(1).title()
    return os.path.splitext(filename)[0].replace("_", " ").title()

def extraer_respuestas_alumno(document):
    respuestas = []

    # 1. Desde tablas
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                letra = cell.text.strip().upper()
                if letra in ['A', 'B', 'C', 'D']:
                    respuestas.append(letra)

    # 2. Desde texto subrayado o resaltado con letras al inicio
    for para in document.paragraphs:
        for run in para.runs:
            texto = run.text.strip()
            if texto and (run.underline or run.font.highlight_color):
                match = re.match(r"^\(?([A-Da-d])", texto)
                if match:
                    letra = match.group(1).upper()
                    respuestas.append(letra)

    return respuestas[:len(clave)]

def calcular_puntaje(respuestas_alumno, clave):
    puntaje = sum(1 for r, c in zip(respuestas_alumno, clave) if r.upper() == c.upper())
    return puntaje

# --- Procesamiento ---
clave = []
if clave_file:
    doc_clave = docx.Document(clave_file)
    clave = extraer_clave_de_tabla(doc_clave)
    st.info("🔐 Clave extraída automáticamente: " + " ".join(clave))

if st.button("📊 Calificar todo") and clave and uploaded_files:
    resultados = []
    for file in uploaded_files:
        doc = docx.Document(file)
        nombre = extraer_nombre(doc, file.name)
        respuestas = extraer_respuestas_alumno(doc)
        puntaje = calcular_puntaje(respuestas, clave)
        resultados.append({
            "Alumno": nombre,
            "Correctas": puntaje,
            "Total Preguntas": len(clave),
            "Nota 20": round((puntaje / len(clave)) * 20, 2)
        })

    df_final = pd.DataFrame(resultados)
    st.success("✅ Evaluación completa")
    st.dataframe(df_final)

    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_final.to_excel(writer, index=False, sheet_name="Resultados")
    st.download_button("⬇️ Descargar resultados en Excel", output.getvalue(), file_name="resultados_examenes.xlsx")

elif not clave:
    st.warning("⚠️ Asegúrate de subir un archivo de claves válido antes de calificar.")
